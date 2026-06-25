"""Word 导出工具。

将 Markdown / TXT / HTML 文档渲染后的 HTML 转换为 .docx，
支持标题、段落、列表、表格、代码块、引用、图片等常见元素。
"""
from __future__ import annotations

import base64
import html
import io
import json
import logging
import mimetypes
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from urllib.parse import unquote

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.runtime import resource_path
from app.utils.markdown_utils import process_markdown
from app.utils.pdf_utils import _fix_image_paths_in_html, _launch_chromium

try:
    from playwright.sync_api import sync_playwright
    _PLAYWRIGHT_AVAILABLE = True
except Exception:  # noqa: BLE001
    sync_playwright = None  # type: ignore
    _PLAYWRIGHT_AVAILABLE = False


SUPPORTED_EXPORT_EXTENSIONS = {'.md', '.markdown', '.txt', '.html', '.htm'}
logger = logging.getLogger(__name__)

# Word 中常用中文字体，按优先级尝试
_CJK_FONT_FAMILIES = ['Noto Sans SC', 'PingFang SC', 'Microsoft YaHei', 'SimHei', 'SimSun']


@dataclass
class WordDocument:
    """待导出的 Word 文档。"""

    path: str
    title: str
    markdown: str
    html: str
    is_html: bool = False


def safe_library_file(base_dir: str | os.PathLike[str], rel_path: str) -> Path:
    """解析并校验文档库内的相对文件路径。"""
    if not rel_path or not isinstance(rel_path, str):
        raise ValueError('文件路径不能为空')

    normalized = rel_path.replace('\\', '/').strip('/')
    base = Path(base_dir).resolve()
    target = (base / normalized).resolve()

    if not str(target).startswith(str(base)):
        raise ValueError('非法路径')
    if not target.exists():
        raise FileNotFoundError('文件不存在')
    if not target.is_file():
        raise ValueError('只能导出文件，不能导出文件夹')
    if target.suffix.lower() not in SUPPORTED_EXPORT_EXTENSIONS:
        raise ValueError('只支持导出 .md、.markdown、.txt、.html 文件')

    return target


_MERMAID_JS_PATH = str(resource_path('static/libs/mermaid.min.js'))

# Mermaid 渲染配置（与 PDF 导出保持一致）
_MERMAID_CONFIG = {
    'startOnLoad': False,
    'theme': 'base',
    'themeVariables': {
        'fontFamily': '-apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", "Helvetica Neue", Arial, sans-serif',
        'fontSize': '14px',
        'background': 'transparent',
        'primaryColor': '#dbeafe',
        'primaryTextColor': '#1e3a5f',
        'primaryBorderColor': '#3b82f6',
        'secondaryColor': '#fef3c7',
        'tertiaryColor': '#dcfce7',
        'lineColor': '#64748b',
        'mainBkg': '#dbeafe',
        'secondBkg': '#fef3c7',
        'tertiaryBkg': '#dcfce7',
        'nodeBorder': '#3b82f6',
        'clusterBkg': '#f8fafc',
        'clusterBorder': '#cbd5e1',
        'titleColor': '#1e293b',
        'edgeLabelBackground': '#ffffff',
        'actorBkg': '#dbeafe',
        'actorBorder': '#3b82f6',
        'actorTextColor': '#1e3a5f',
        'signalColor': '#475569',
        'signalTextColor': '#1e293b',
        'noteBkgColor': '#fef9c3',
        'noteTextColor': '#713f12',
        'noteBorderColor': '#facc15',
        'pie1': '#3b82f6', 'pie2': '#10b981', 'pie3': '#f59e0b', 'pie4': '#ec4899',
        'pie5': '#8b5cf6', 'pie6': '#06b6d4', 'pie7': '#f43f5e', 'pie8': '#84cc16',
        'pieStrokeColor': '#ffffff',
        'pieOuterStrokeColor': '#cbd5e1',
        'sectionBkgColor': '#dbeafe',
        'altSectionBkgColor': '#fef3c7',
        'gridColor': '#e2e8f0',
        'taskBkgColor': '#3b82f6',
        'taskTextColor': '#ffffff',
    },
    'securityLevel': 'loose',
    'suppressErrorRendering': False,
    'fontFamily': '-apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", "Helvetica Neue", Arial, sans-serif',
    'logLevel': 'error',
    'flowchart': {'useMaxWidth': True, 'htmlLabels': True, 'curve': 'basis', 'padding': 18, 'nodeSpacing': 55, 'rankSpacing': 60},
    'sequence': {'useMaxWidth': True, 'showSequenceNumbers': False, 'wrap': True},
    'class': {'useMaxWidth': True},
    'state': {'useMaxWidth': True},
    'er': {'useMaxWidth': True},
    'gantt': {'useMaxWidth': True, 'barHeight': 22, 'barGap': 6},
    'pie': {'useMaxWidth': True},
    'journey': {'useMaxWidth': True},
    'gitGraph': {'useMaxWidth': True},
    'mindmap': {'useMaxWidth': True, 'padding': 12},
    'timeline': {'useMaxWidth': True, 'padding': 12},
    'sankey': {'useMaxWidth': True},
    'xyChart': {'useMaxWidth': True},
    'quadrantChart': {'useMaxWidth': True},
    'requirement': {'useMaxWidth': True},
    'block': {'useMaxWidth': True, 'padding': 12},
    'packet': {'useMaxWidth': True, 'padding': 12},
    'architecture': {'useMaxWidth': True, 'padding': 12},
    'kanban': {'useMaxWidth': True, 'padding': 12},
}


def _capture_mermaid_screenshots(browser, html_fragment: str, block_count: int) -> list[bytes]:
    """在已启动的浏览器实例中，将 HTML 片段中的 Mermaid 代码块渲染为 PNG 截图。

    返回与代码块顺序对应的 PNG 字节列表。
    """
    from playwright.sync_api import Error as PlaywrightError

    render_html = f'''<!doctype html>
<html lang="zh-CN">
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        body {{ margin: 20px; background: #ffffff; color: #1f2937; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", "Helvetica Neue", Arial, sans-serif; }}
        .markdown-body {{ max-width: 900px; margin: 0 auto; }}
        .mermaid {{ margin: 16px 0; }}
    </style>
</head>
<body>
    <div class="markdown-body">{html_fragment}</div>
    <script>
    window.__MARKINOTE_MERMAID_READY__ = false;
    window.__MARKINOTE_MERMAID_COUNT__ = {block_count};
    </script>
</body>
</html>'''

    page = browser.new_page(viewport={'width': 1280, 'height': 1800}, device_scale_factor=2)
    try:
        page.set_content(render_html, wait_until='networkidle', timeout=60000)
        page.add_script_tag(path=_MERMAID_JS_PATH)
        page.wait_for_function('typeof window.mermaid !== "undefined"', timeout=30000)
        mermaid_config_json = json.dumps(_MERMAID_CONFIG, ensure_ascii=False)
        page.evaluate(
            f'''async () => {{
                mermaid.initialize({mermaid_config_json});

                const blocks = Array.from(document.querySelectorAll('pre code.language-mermaid'));
                window.__MARKINOTE_MERMAID_COUNT__ = blocks.length;
                for (let i = 0; i < blocks.length; i++) {{
                    const codeBlock = blocks[i];
                    const pre = codeBlock.parentElement;
                    const source = codeBlock.textContent.trim();
                    const container = document.createElement('div');
                    container.className = 'mermaid';
                    container.id = 'markinote-mermaid-' + i;
                    pre.parentNode.replaceChild(container, pre);
                    try {{
                        const result = await mermaid.render('mermaid-' + i, source);
                        container.innerHTML = result.svg;
                    }} catch (error) {{
                        container.innerHTML = '<pre style="color:#b91c1c; white-space:pre-wrap;">Mermaid 渲染失败：' + (error.message || error) + '</pre>';
                    }}
                }}
                if (document.fonts && document.fonts.ready) await document.fonts.ready;
                window.__MARKINOTE_MERMAID_READY__ = true;
            }}'''
        )
        page.wait_for_function(
            'window.__MARKINOTE_MERMAID_READY__ === true',
            timeout=int(os.environ.get('MARKINOTE_PDF_RENDER_TIMEOUT_MS', '60000')),
        )
        elements = page.query_selector_all('.mermaid')
        return [el.screenshot(type='png') for el in elements]
    finally:
        page.close()


def _render_mermaid_blocks(html_content: str, browser=None) -> str:
    """将 HTML 中的 Mermaid 代码块渲染为图片并替换为 <img>。

    如果传入 browser，则复用该浏览器实例以节省启动时间。
    如果系统没有可用的 Chromium/Chrome/Edge，则保留源代码不阻断导出。
    """
    if 'language-mermaid' not in html_content:
        return html_content

    soup = BeautifulSoup(html_content, 'html.parser')
    mermaid_pres: list[Tag] = []
    for pre in soup.find_all('pre'):
        code = pre.find('code')
        if code and 'language-mermaid' in ' '.join(code.get('class', [])):
            mermaid_pres.append(pre)

    if not mermaid_pres:
        return html_content

    try:
        if browser is None:
            with sync_playwright() as playwright:
                browser = _launch_chromium(playwright)
                try:
                    screenshots = _capture_mermaid_screenshots(browser, str(soup), len(mermaid_pres))
                finally:
                    browser.close()
        else:
            screenshots = _capture_mermaid_screenshots(browser, str(soup), len(mermaid_pres))
    except Exception as exc:  # noqa: BLE001 - 渲染失败应降级为源代码
        logger.warning('Mermaid 渲染失败，Word 导出将保留源代码: %s', exc)
        return html_content

    for idx, pre in enumerate(mermaid_pres):
        if idx >= len(screenshots) or not screenshots[idx]:
            continue
        img = soup.new_tag('img')
        img['src'] = f'data:image/png;base64,{base64.b64encode(screenshots[idx]).decode("ascii")}'
        img['alt'] = 'Mermaid diagram'
        img['style'] = 'max-width: 100%;'
        pre.replace_with(img)

    return str(soup)


def read_word_document(base_dir: str | os.PathLike[str], rel_path: str) -> WordDocument:
    """读取单个文档并渲染为 HTML（图片已内嵌为 data URI）。"""
    target = safe_library_file(base_dir, rel_path)
    raw = target.read_text(encoding='utf-8')
    is_html = target.suffix.lower() in ('.html', '.htm')
    # HTML 文件直接使用原文作为渲染结果，不走 Markdown 渲染管道
    rendered = raw if is_html else process_markdown(raw)
    # 复用 PDF 工具将图片相对路径转为 data URI，确保图片可被安全内嵌
    rendered = _fix_image_paths_in_html(rendered, rel_path.replace('\\', '/').strip('/'), base_dir)
    return WordDocument(
        path=rel_path.replace('\\', '/').strip('/'),
        title=target.name,
        markdown=raw,
        html=rendered,
        is_html=is_html,
    )


def slugify_filename(name: str, default: str = 'MarkiNote') -> str:
    """生成适合下载的 Word 文件名。"""
    stem = Path(name).stem or default
    stem = re.sub(r'[\\/:*?"<>|\x00-\x1f]+', '_', stem).strip(' ._')
    return f'{stem or default}.docx'


def _set_run_font(run, font_name: str | None = None, size_pt: float | None = None) -> None:
    """设置 run 的字体与字号。"""
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _set_paragraph_cjk_font(paragraph, font_name: str | None = None, size_pt: float | None = None) -> None:
    """设置段落中所有 run 的字体。"""
    for run in paragraph.runs:
        _set_run_font(run, font_name, size_pt)


def _hex_to_rgb(hex_color: str) -> RGBColor | None:
    """将 #RRGGBB 转为 RGBColor。"""
    hex_color = hex_color.lstrip('#')
    if len(hex_color) != 6:
        return None
    try:
        return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))
    except ValueError:
        return None


def _set_cell_shading(cell, color: str) -> None:
    """设置单元格背景色（#RRGGBB）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color.lstrip('#'))
    tcPr.append(shd)


def _remove_table_borders(table) -> None:
    """移除表格所有边框。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _get_text(element: Tag | NavigableString) -> str:
    """提取元素的纯文本，保留换行。"""
    if isinstance(element, NavigableString):
        return str(element)
    return element.get_text(separator='\n', strip=True)


def _add_inline_contents(
    paragraph,
    element: Tag | NavigableString,
    base_dir: str | os.PathLike[str] = '',
    doc_path: str = '',
) -> None:
    """将元素内的行内内容添加到段落中。

    如果 element 自身就是 strong/em/code/a 等行内标签，会保留对应格式；
    否则会递归处理其子元素。
    """
    if isinstance(element, NavigableString):
        text = str(element)
        if text:
            run = paragraph.add_run(text)
            _set_run_font(run, size_pt=11)
        return

    tag = element.name
    if tag == 'br':
        paragraph.add_run('\n')
        return
    if tag == 'img':
        # 图片作为独立段落处理更稳定；此处先添加占位说明，
        # 外层块级转换逻辑会单独捕获 img 标签。
        alt = element.get('alt', '') or ''
        paragraph.add_run(f'[图片: {alt}]')
        return
    if tag in ('strong', 'b'):
        text = _get_text(element)
        if text:
            run = paragraph.add_run(text)
            run.bold = True
            _set_run_font(run, size_pt=11)
        return
    if tag in ('em', 'i'):
        text = _get_text(element)
        if text:
            run = paragraph.add_run(text)
            run.italic = True
            _set_run_font(run, size_pt=11)
        return
    if tag == 'del':
        text = _get_text(element)
        if text:
            run = paragraph.add_run(text)
            run.font.strike = True
            _set_run_font(run, size_pt=11)
        return
    if tag == 'code':
        text = _get_text(element)
        if text:
            run = paragraph.add_run(text)
            run.font.name = 'Courier New'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
            run.font.size = Pt(10)
        return
    if tag == 'a':
        text = _get_text(element)
        href = element.get('href', '')
        if text or href:
            run = paragraph.add_run(text or href)
            run.font.color.rgb = _hex_to_rgb('2563EB') or RGBColor(0x25, 0x63, 0xEB)
            run.font.underline = True
            _set_run_font(run, size_pt=11)
        return
    if tag in ('span', 'sub', 'sup'):
        # 数学公式等 span 保留原始文本
        text = _get_text(element)
        if text:
            run = paragraph.add_run(text)
            _set_run_font(run, size_pt=11)
        return

    # 递归处理子元素
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                _set_run_font(run, size_pt=11)
            continue
        _add_inline_contents(paragraph, child, base_dir, doc_path)


def _add_image(
    document: Document,
    img_tag: Tag,
    base_dir: str | os.PathLike[str] = '',
    doc_path: str = '',
) -> None:
    """将图片插入到 Word 文档。"""
    src = img_tag.get('src', '')
    alt = img_tag.get('alt', '') or ''
    if not src:
        return

    image_stream: io.BytesIO | None = None
    image_ext = '.png'

    if src.startswith('data:'):
        try:
            header, encoded = src.split(',', 1)
            mime = header.split(';')[0].split(':')[1] if ':' in header else 'image/png'
            image_ext = mimetypes.guess_extension(mime) or '.png'
            image_stream = io.BytesIO(base64.b64decode(encoded))
        except Exception:  # noqa: BLE001
            logger.debug('Failed to decode data URI image', exc_info=True)
    elif src.startswith(('http://', 'https://')):
        # 外部 URL 暂时以链接文本替代
        p = document.add_paragraph()
        run = p.add_run(f'[图片链接: {alt or src}]')
        run.font.color.rgb = _hex_to_rgb('2563EB') or RGBColor(0x25, 0x63, 0xEB)
        run.font.underline = True
        _set_run_font(run, size_pt=11)
        return
    else:
        # 相对路径：拼接为文档库内绝对路径
        doc_dir = doc_path.rsplit('/', 1)[0] if '/' in doc_path else ''
        library_base = Path(base_dir).resolve()
        image_rel = f'{doc_dir}/{src}' if doc_dir else src
        image_abs = (library_base / image_rel).resolve()
        if str(image_abs).startswith(str(library_base)) and image_abs.is_file():
            image_stream = io.BytesIO(image_abs.read_bytes())
            image_ext = image_abs.suffix or '.png'

    if image_stream is None:
        p = document.add_paragraph()
        run = p.add_run(f'[图片: {alt or src}]')
        _set_run_font(run, size_pt=11)
        return

    try:
        # 限制图片宽度为页面可阅读区域（约 6 英寸）
        picture = document.add_picture(image_stream, width=Inches(5.8))
        # 居中对齐
        picture_paragraph = picture._element.getparent()
        if picture_paragraph is not None:
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:  # noqa: BLE001
        logger.debug('Failed to embed image into Word', exc_info=True)
        p = document.add_paragraph()
        run = p.add_run(f'[图片: {alt or src}]')
        _set_run_font(run, size_pt=11)


def _add_code_block(document: Document, pre_element: Tag) -> None:
    """将 <pre> 代码块作为带灰色背景的表格单元格插入。"""
    code_tag = pre_element.find('code')
    if code_tag:
        # 使用默认分隔符，避免 codehilite 高亮生成的 <span> 之间被插入多余换行
        code_text = code_tag.get_text()
    else:
        code_text = pre_element.get_text()

    # 去掉首尾多余空行
    code_text = code_text.strip('\n')
    if not code_text:
        return

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, 'F3F4F6')

    # 单元格默认已有一个段落，先清空
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.clear()
    cell_paragraph.paragraph_format.space_before = Pt(6)
    cell_paragraph.paragraph_format.space_after = Pt(6)
    cell_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for line in code_text.split('\n'):
        run = cell_paragraph.add_run(line)
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
        run.font.size = Pt(9.5)
        cell_paragraph.add_run('\n')

    # 移除尾部多余换行 run
    if cell_paragraph.runs and cell_paragraph.runs[-1].text == '\n':
        cell_paragraph.runs[-1]._element.getparent().remove(cell_paragraph.runs[-1]._element)


def _add_table(document: Document, table_element: Tag) -> None:
    """将 HTML 表格转换为 Word 表格。"""
    rows = table_element.find_all('tr')
    if not rows:
        return

    col_count = max(len(row.find_all(['td', 'th'])) for row in rows)
    if col_count == 0:
        return

    docx_table = document.add_table(rows=len(rows), cols=col_count)
    docx_table.style = 'Table Grid'
    docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells):
            if col_idx >= col_count:
                break
            docx_cell = docx_table.rows[row_idx].cells[col_idx]
            docx_cell.text = ''
            # 表头加粗
            is_header = cell.name == 'th' or row.parent.name == 'thead'
            p = docx_cell.paragraphs[0]
            p.clear()
            _add_inline_contents(p, cell)
            if is_header:
                for run in p.runs:
                    run.bold = True


def _add_list(
    document: Document,
    list_element: Tag,
    base_dir: str | os.PathLike[str] = '',
    doc_path: str = '',
    level: int = 0,
) -> None:
    """将 HTML 列表转换为 Word 列表段落。"""
    is_ordered = list_element.name == 'ol'
    style_name = 'List Number' if is_ordered else 'List Bullet'

    def _new_list_paragraph(first: bool = False) -> object:
        p = document.add_paragraph(style=style_name if first else None)
        # 缩进模拟嵌套层级（python-docx 的 List Bullet 样式对嵌套支持有限）
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        p.paragraph_format.first_line_indent = Inches(-0.25)
        return p

    for li in list_element.find_all('li', recursive=False):
        p = _new_list_paragraph(first=True)
        first = False

        for child in li.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    run = p.add_run(text)
                    _set_run_font(run, size_pt=11)
                continue

            if not isinstance(child, Tag):
                continue

            if child.name in ('ul', 'ol'):
                # 结束当前段落，处理嵌套列表
                _add_list(document, child, base_dir, doc_path, level + 1)
                # 嵌套列表后的内容需要新段落
                p = _new_list_paragraph()
            elif child.name in ('p', 'div'):
                # 段落分隔：先结束当前段落（若已有内容），再创建新段落承载块级内容
                if p.text.strip() or p.runs:
                    p = _new_list_paragraph()
                _add_inline_contents(p, child, base_dir, doc_path)
            else:
                _add_inline_contents(p, child, base_dir, doc_path)


def _convert_block_element(
    document: Document,
    element: Tag,
    base_dir: str | os.PathLike[str] = '',
    doc_path: str = '',
) -> None:
    """转换单个块级元素。"""
    tag = element.name

    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        heading = document.add_heading(level=min(level, 9))
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.add_run(_get_text(element))
        _set_run_font(run, size_pt=16 - level + 1 if 16 - level + 1 >= 10 else 10)
    elif tag == 'p':
        # 检查是否是占位符段落（数学公式/Mermaid 已被恢复为具体标签）
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _add_inline_contents(p, element, base_dir, doc_path)
        if not p.text.strip() and not p.runs:
            # 空段落不保留
            p._element.getparent().remove(p._element)
    elif tag == 'blockquote':
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(6)
        _add_inline_contents(p, element, base_dir, doc_path)
    elif tag in ('ul', 'ol'):
        _add_list(document, element, base_dir, doc_path)
    elif tag == 'pre':
        _add_code_block(document, element)
    elif tag == 'table':
        _add_table(document, element)
    elif tag == 'hr':
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('─' * 40)
        run.font.color.rgb = _hex_to_rgb('9CA3AF') or RGBColor(0x9C, 0xA3, 0xAF)
        _set_run_font(run, size_pt=11)
    elif tag == 'img':
        _add_image(document, element, base_dir, doc_path)
    elif tag == 'div':
        # math-block, highlight 等 div 容器
        cls = ' '.join(element.get('class', []))
        if 'math-block' in cls or 'math-inline' in cls:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(_get_text(element))
            run.font.name = 'Cambria Math'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Cambria Math')
            _set_run_font(run, size_pt=11)
        elif 'highlight' in cls:
            # 代码高亮容器，内部通常有 <pre>
            pre = element.find('pre')
            if pre:
                _add_code_block(document, pre)
            else:
                for child in element.children:
                    if isinstance(child, Tag):
                        _convert_block_element(document, child, base_dir, doc_path)
        else:
            for child in element.children:
                if isinstance(child, Tag):
                    _convert_block_element(document, child, base_dir, doc_path)
                elif isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        p = document.add_paragraph(text)
                        _set_paragraph_cjk_font(p, size_pt=11)
    elif tag in ('script', 'style', 'noscript'):
        # 忽略脚本与样式标签
        return
    else:
        # 其他未识别标签：尝试提取文本
        text = _get_text(element)
        if text:
            p = document.add_paragraph(text)
            _set_paragraph_cjk_font(p, size_pt=11)


def _add_html_to_document(
    document: Document,
    html_content: str,
    base_dir: str | os.PathLike[str] = '',
    doc_path: str = '',
) -> None:
    """将 HTML 内容追加到已有的 Word Document 中。"""
    soup = BeautifulSoup(html_content, 'html.parser')
    root = soup.body or soup

    # 优先使用 .markdown-body 容器；HTML 原文则可能只有 body
    container = root.find(class_='markdown-body') or root

    for child in container.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                p = document.add_paragraph(text)
                _set_paragraph_cjk_font(p, size_pt=11)
        elif isinstance(child, Tag):
            _convert_block_element(document, child, base_dir, doc_path)


def convert_html_to_docx(
    html_content: str,
    *,
    title: str = 'MarkiNote Export',
    base_dir: str | os.PathLike[str] = '',
    doc_path: str = '',
) -> Document:
    """将 HTML 内容转换为 python-docx Document 对象。"""
    document = Document()

    # 设置默认正文字体
    style = document.styles['Normal']
    style.font.name = 'Noto Sans SC'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans SC')
    style.font.size = Pt(11)

    # 设置页面边距（A4 默认左右约 2.54cm，这里调窄一些）
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    _add_html_to_document(document, html_content, base_dir, doc_path)
    return document


def render_word_bytes(
    documents: Iterable[WordDocument],
    *,
    title: str = 'MarkiNote Export',
) -> bytes:
    """将一份或多份文档渲染为 Word 字节。"""
    docs = list(documents)
    if not docs:
        raise ValueError('请至少选择一个文档')

    # 批量渲染 Mermaid 图表：多文档共享一次浏览器实例，避免重复启动
    playwright_obj = None
    browser = None
    try:
        if _PLAYWRIGHT_AVAILABLE and any('language-mermaid' in doc.html for doc in docs):
            playwright_obj = sync_playwright().start()
            browser = _launch_chromium(playwright_obj)
            for doc in docs:
                if 'language-mermaid' in doc.html:
                    doc.html = _render_mermaid_blocks(doc.html, browser=browser)
    except Exception as exc:  # noqa: BLE001 - 渲染失败应降级为源代码
        logger.warning('Mermaid 批量渲染失败，Word 导出将保留源代码: %s', exc)
    finally:
        if browser:
            browser.close()
        if playwright_obj:
            playwright_obj.stop()

    if len(docs) == 1:
        doc = docs[0]
        document = convert_html_to_docx(doc.html, title=doc.title, doc_path=doc.path)
    else:
        document = Document()
        # 设置默认字体与页面边距
        style = document.styles['Normal']
        style.font.name = 'Noto Sans SC'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans SC')
        style.font.size = Pt(11)
        for section in document.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # 汇总标题
        heading = document.add_heading(level=1)
        run = heading.add_run(title)
        _set_run_font(run, size_pt=20)
        document.add_paragraph()

        for index, doc in enumerate(docs):
            if index > 0:
                document.add_page_break()
            doc_heading = document.add_heading(level=1)
            run = doc_heading.add_run(doc.title)
            _set_run_font(run, size_pt=18)
            path_para = document.add_paragraph(doc.path)
            path_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in path_para.runs:
                run.font.color.rgb = _hex_to_rgb('6B7280') or RGBColor(0x6B, 0x72, 0x80)
                run.font.size = Pt(10)
            document.add_paragraph()

            _add_html_to_document(document, doc.html, doc_path=doc.path)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()
